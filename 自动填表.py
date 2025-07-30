from datetime import datetime, timedelta
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

# 定义起止日期和总周数，需要自己修改
start_date = datetime(2025, 2, 17) 
total_weeks = 18

# 读取模板文件路径和保存路径
template_file = r".\xxx宿舍xx学院第1周xxxx级xxxx班学xx号楼xxx宿舍回寝记录表.docx"
save_path = r'.\安全表格'

# 批量生成18周的考勤表
current_start_date = start_date

for week in range(1, total_weeks + 1):
    # 读取模板
    doc = Document(template_file)

    # 清空表格第一行第一个单元格
    table = doc.tables[0]
    table.cell(0, 0).text = '' 

    # 重新写入格式化标题
    new_title = f'第 {week} 周   202x 级 xxxx xx 班学 xx 号楼 xxx 宿舍考勤表'
    p = table.cell(0, 0).paragraphs[0]
    run = p.add_run(new_title)

    # 设置格式
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run.font.name = '宋体'  # 字体为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 兼容中文宋体设置
    run.font.size = Pt(18)  # 字号18
    
    # 给“第 x 周”中的数字部分加下划线
    start_idx = new_title.find(f'{week}')
    end_idx = start_idx + len(str(week))
    run.clear()
    
    # 重写，拆分加下划线的部分
    run1 = p.add_run(new_title[:start_idx])
    run1.font.name = '宋体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run1.font.size = Pt(18)
    
    run2 = p.add_run(new_title[start_idx:end_idx])
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(18)
    run2.font.underline = True
    
    run3 = p.add_run(new_title[end_idx:])
    run3.font.name = '宋体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run3.font.size = Pt(18)

    # 填充日期（跳过第一行，从第二行的第一列开始填充）
    for i in range(7):
        date = (current_start_date + timedelta(days=i)).strftime('%Y/%m/%d')
        table.cell(i + 2, 0).text = date
    
    # 保存文件
    file_name = f'xxx宿舍xx学院第{week}周xxxx级xxxx班学xx号楼xxx宿舍回寝记录表.docx'
    file_path = f'{save_path}\\{file_name}'
    doc.save(file_path)
    
    print(f'已生成：{file_path}')
    
    # 更新下一周的开始日期
    current_start_date += timedelta(weeks=1)

print("所有考勤表生成完毕！")